import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def get_items_list(input_file):
    doc = Document(input_file)
    items = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            items.append(text)
    return items

def apply_golden_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    
    for border_name in borders:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D4AF37')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def add_hotel_header(cell):
    cell.text = ''
    
    logo_para = cell.paragraphs[0]
    logo_run = logo_para.add_run('ðŸ¨\n')
    logo_run.font.size = Pt(24)
    logo_run.font.color.rgb = RGBColor(139, 195, 74)
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name_para = cell.add_paragraph()
    name_run = name_para.add_run('NARAYANI HEIGHTS')
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = 'Arial'
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle_para = cell.add_paragraph()
    subtitle_run = subtitle_para.add_run('HOTEL AND RESORT')
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.name = 'Arial'
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    stars_para = cell.add_paragraph()
    stars_run = stars_para.add_run('â­ â­ â­ â­')
    stars_run.font.size = Pt(10)
    stars_run.font.color.rgb = RGBColor(255, 215, 0)
    stars_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    web_para = cell.add_paragraph()
    web_run = web_para.add_run('WWW.NARAYANIHEIGHTS.COM')
    web_run.font.size = Pt(9)
    web_run.font.name = 'Arial'
    web_run.bold = True
    web_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_food_item_template(doc, item_name, item_number):
    main_table = doc.add_table(rows=1, cols=2)
    main_table.autofit = False
    main_table.allow_autofit = False
    
    main_table.columns[0].width = Inches(2.5)
    main_table.columns[1].width = Inches(4.0)
    
    left_cell = main_table.rows[0].cells[0]
    right_cell = main_table.rows[0].cells[1]
    
    add_hotel_header(left_cell)
    
    tcPr = left_cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    right_border = OxmlElement('w:right')
    right_border.set(qn('w:val'), 'single')
    right_border.set(qn('w:sz'), '12')
    right_border.set(qn('w:space'), '0')
    right_border.set(qn('w:color'), '000000')
    tcBorders.append(right_border)
    tcPr.append(tcBorders)
    
    right_cell.text = ''
    
    item_words = item_name.upper().split()
    
    for word in item_words:
        p = right_cell.add_paragraph() if item_words.index(word) > 0 else right_cell.paragraphs[0]
        r = p.add_run(word)
        r.bold = True
        r.font.size = Pt(48)
        r.font.name = 'Playfair Display'
        r.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(0)
    
    right_cell.paragraphs[0].space_before = Pt(60)
    right_cell.paragraphs[-1].space_after = Pt(60)
    
    apply_golden_border(main_table)
    
    doc.add_paragraph()

def create_formatted_docx(items, output_file):
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    for i, item in enumerate(items, start=1):
        add_food_item_template(doc, item, i)
        
        if i % 2 == 0 and i < len(items):
            doc.add_page_break()
        elif i % 2 == 1 and i < len(items):
            sep_para = doc.add_paragraph()
            sep_para.add_run('\n')
            doc.add_paragraph()
    
    doc.save(output_file)
    print(f"âœ“ Document created successfully: {output_file}")
    print(f"âœ“ Total items processed: {len(items)}")
    print(f"âœ“ Total pages: {(len(items) + 1) // 2 + 1}")

def main():
    print("=" * 70)
    print("Food Items Document Processor")
    print("=" * 70)
    
    input_file = input("\nEnter the path to input DOCX file (with items list): ").strip('"')
    
    if not os.path.exists(input_file):
        print(f"âœ— Error: File '{input_file}' not found!")
        return
    
    output_file = input("Enter the output file name (default: food_items_formatted.docx): ").strip('"')
    if not output_file:
        output_file = "food_items_formatted.docx"
    
    if not output_file.endswith('.docx'):
        output_file += '.docx'
    
    print("\nProcessing...")
    
    try:
        items = get_items_list(input_file)
        
        if not items:
            print("âœ— No items found in the input file!")
            return
        
        print(f"âœ“ Found {len(items)} items in input file")
        
        create_formatted_docx(items, output_file)
        
    except Exception as e:
        print(f"\nâœ— Error occurred: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
