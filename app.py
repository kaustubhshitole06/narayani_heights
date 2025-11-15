import os
import shutil
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"

if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyAgSwpTwsqp49Pr-oD6uP88ZpXFgXMjmPo")
genai.configure(api_key=GEMINI_API_KEY)

def get_items_list(input_file):
    doc = Document(input_file)
    items = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            items.append(text)
    return items

def extract_food_items_from_pdf(pdf_path):
    try:
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        pdf_file = genai.upload_file(pdf_path)
        
        prompt = """Extract ONLY the food item names from this PDF menu. 
        Return each item on a new line.
        Do not include prices, descriptions, categories, or any other text.
        Just the food item names, one per line.
        If there are no food items, return empty."""
        
        response = model.generate_content([prompt, pdf_file])
        
        genai.delete_file(pdf_file.name)
        
        if response.text:
            items = [item.strip() for item in response.text.strip().split('\n') if item.strip()]
            return items
        
        return []
    except Exception as e:
        raise Exception(f"Failed to extract items from PDF: {str(e)}")

def apply_golden_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '24')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D4AF37')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def add_hotel_header(cell):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run('🏨\n')
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(139, 195, 74)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = cell.add_paragraph()
    r = p.add_run('NARAYANI HEIGHTS')
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = 'Arial'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = cell.add_paragraph()
    r = p.add_run('HOTEL AND RESORT')
    r.font.size = Pt(10)
    r.font.name = 'Arial'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = cell.add_paragraph()
    r = p.add_run('⭐ ⭐ ⭐ ⭐')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(255, 215, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = cell.add_paragraph()
    r = p.add_run('WWW.NARAYANIHEIGHTS.COM')
    r.font.size = Pt(9)
    r.font.name = 'Arial'
    r.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_food_item_template(doc, item_name, item_number):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.0)
    
    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    
    add_hotel_header(left)
    
    tcPr = left._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    rb = OxmlElement('w:right')
    rb.set(qn('w:val'), 'single')
    rb.set(qn('w:sz'), '12')
    rb.set(qn('w:space'), '0')
    rb.set(qn('w:color'), '000000')
    tcBorders.append(rb)
    tcPr.append(tcBorders)
    
    right.text = ''
    p = right.paragraphs[0]
    r = p.add_run(item_name)
    r.bold = True
    r.font.size = Pt(36)
    r.font.name = 'Arial'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(80)
    p.space_after = Pt(80)
    
    apply_golden_border(table)
    doc.add_paragraph()

def create_formatted_docx(items, output_file):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    title = doc.add_heading('Food Items Catalog', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph('Complete Item Information')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].italic = True
    
    doc.add_paragraph('\n' * 3)
    doc.add_page_break()
    
    for i, item in enumerate(items, 1):
        add_food_item_template(doc, item, i)
        if i % 2 == 0 and i < len(items):
            doc.add_page_break()
        elif i % 2 == 1 and i < len(items):
            doc.add_paragraph('\n' + '═' * 70 + '\n')
    
    doc.save(output_file)

@app.post("/process")
async def process_file(file: UploadFile = File(...)):
    if not (file.filename.endswith('.docx') or file.filename.endswith('.pdf')):
        raise HTTPException(400, "Only .docx and .pdf files allowed")
    
    input_path = os.path.join(UPLOAD_DIR, file.filename)
    output_name = f"formatted_{file.filename.rsplit('.', 1)[0]}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_name)
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
        
        if file.filename.endswith('.pdf'):
            items = extract_food_items_from_pdf(input_path)
        else:
            items = get_items_list(input_path)
            
        if not items:
            raise HTTPException(400, "No items found")
        
        create_formatted_docx(items, output_path)
        
        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=output_name
        )
    except Exception as e:
        raise HTTPException(500, str(e))
    finally:
        if os.path.exists(input_path):
            os.remove(input_path)

app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
